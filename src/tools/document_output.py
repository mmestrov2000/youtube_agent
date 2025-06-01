import os
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import boto3
from dotenv import load_dotenv
from agno.tools import tool
from typing import Annotated
import re

class ReportProcessor:
    def __init__(self):
        load_dotenv()
        self.api_key = os.getenv("OPENAI_API_KEY")
        self.s3_bucket = os.getenv("S3_BUCKET_NAME")
        self.client = OpenAI(api_key=self.api_key)
        self.s3 = boto3.client(
            "s3",
            aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
            region_name="eu-north-1",
            config=boto3.session.Config(
                signature_version="s3v4",
                region_name="eu-north-1"
            )
        )

    def extract_structure(self, report_text):
        # Sanitize the input text to remove non-XML compatible characters
        def sanitize_text(text):
            # Remove NULL bytes and control characters
            text = ''.join(char for char in text if ord(char) >= 32 or char in "\n\r\t")
            # Replace any remaining problematic characters with spaces
            text = "".join(char if ord(char) < 0x10000 else " " for char in text)
            return text

        sanitized_text = sanitize_text(report_text)

        system_prompt = """
        You are an expert at document structuring. Given an unstructured or semi-structured report,
        extract the following in a dictionary format:
        {
          "title": "Document Title",
          "sections": [
            {
              "heading": "Section Heading",
              "content": [
                "Bullet point 1 or paragraph",
                "Bullet point 2 or paragraph"
              ]
            },
            ...
          ]
        }
        Keep the content concise and usable in a Word document.
        """

        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": sanitized_text},
            ],
        )
        return eval(response.choices[0].message.content)

    def create_document(self, structure_dict, output_path="structured_report.docx"):
        doc = Document()

        # Add title
        title = doc.add_heading(structure_dict["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        for section in structure_dict["sections"]:
            doc.add_heading(section["heading"], level=1)
            for content in section["content"]:
                doc.add_paragraph(
                    content,
                    style="List Bullet" if content.startswith("*") else "Normal",
                )

        # Save locally
        doc.save(output_path)

        # Upload to S3
        s3_key = os.path.basename(output_path)
        self.s3.upload_file(output_path, self.s3_bucket, s3_key)

        # Delete local file after successful upload
        if os.path.exists(output_path):
            os.remove(output_path)

        # Generate presigned URL
        url = self.s3.generate_presigned_url(
            "get_object",
            Params={
                "Bucket": self.s3_bucket,
                "Key": s3_key,
            },
            ExpiresIn=3600,
        )

        return url

@tool(
    name="generate_structured_document",
    description="Generate a structured .docx document from user-provided content and return a URL to the file.",
    show_result=True,
    cache_results=True,
    cache_ttl=1800
)
def Document_Output(report_text: Annotated[str, 'Provide the full content required to generate the document']):
    processor = ReportProcessor()
    structured_data = processor.extract_structure(report_text)
    
    # Generate filename from title
    title = structured_data["title"]
    # Remove special characters and replace spaces with underscores
    safe_title = re.sub(r'[^\w\s-]', '', title)
    safe_title = re.sub(r'[-\s]+', '_', safe_title).strip('-_')
    filename = f"{safe_title}.docx"
    
    url = processor.create_document(structured_data, filename)
    return {'result': f"url of report -> {url}"}